import io
try:
    import pymupdf as fitz
except ImportError:
    import fitz
import docx
import httpx
from bs4 import BeautifulSoup


def extract_pdf(file_bytes: bytes) -> str:
    text_chunks = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text = page.get_text()
            if text.strip():
                text_chunks.append(text)
    full_text = "\n".join(text_chunks).strip()
    if not full_text:
        raise ValueError("Scanned or image-only PDF with no extractable text layer is not supported in v1")
    return full_text


def extract_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs).strip()


def extract_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1").strip()


def extract_url(url: str) -> str:
    with httpx.Client(timeout=10.0, follow_redirects=True) as client:
        response = client.get(url)
        response.raise_for_status()
        html = response.text

    soup = BeautifulSoup(html, "html.parser")
    for element in soup(["script", "style", "nav", "footer", "header"]):
        element.decompose()
    text = soup.get_text(separator="\n")
    lines = (line.strip() for line in text.splitlines())
    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
    return "\n".join(chunk for chunk in chunks if chunk).strip()


def extract(source_input: bytes | str, file_type: str) -> str:
    file_type = file_type.lower()
    if file_type == "url":
        if isinstance(source_input, bytes):
            source_input = source_input.decode("utf-8")
        return extract_url(source_input)
    elif isinstance(source_input, bytes):
        if file_type == "pdf":
            return extract_pdf(source_input)
        elif file_type == "docx":
            return extract_docx(source_input)
        elif file_type == "txt":
            return extract_txt(source_input)
    raise ValueError(f"Unsupported file type or invalid input for type: {file_type}")
